from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import docx
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx2pdf import convert
import os
import json
import base64
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime, date
import psycopg2
from psycopg2.extras import RealDictCursor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from lxml import etree

# Importar la configuración de la base de datos desde etl_app/config.py
try:
    from config import DATABASE_CONFIG as DB_CONFIG
except ImportError:
    # Fallback para desarrollo/testing si se ejecuta directamente el módulo
    DB_CONFIG = {
        "host": "localhost",
        "database": "jupe",
        "user": "postgres",
        "password": "12345678",
        "port": 5432
    }
    logging.warning("No se pudo importar DATABASE_CONFIG desde config.py. Usando configuración por defecto.")

router = APIRouter()
logger = logging.getLogger(__name__)

class DocumentGenerationRequest(BaseModel):
    cedula: str
    tipo_solicitud: str

def get_db_connection():
    """Crear conexión a la base de datos"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        logger.error(f"Error conectando a la base de datos: {e}")
        raise HTTPException(status_code=500, detail="Error de conexión a la base de datos")

def get_funcionario_data(cedula: str) -> Optional[Dict]:
    """Obtiene los datos del funcionario y sus expedientes de la base de datos."""
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)

        logger.info(f"[RECORD_GEN] Buscando funcionario con cédula: {cedula}")

        # Obtener datos del funcionario
        cursor.execute("""
            SELECT 
                f.id,
                f.cedula, 
                f.nombre_completo, 
                f.rango_actual
            FROM funcionarios f
            WHERE f.cedula = %s
        """, (cedula,))
        funcionario = cursor.fetchone()

        if not funcionario:
            logger.warning(f"[RECORD_GEN] Funcionario con cédula {cedula} no encontrado")
            return None

        logger.info(f"[RECORD_GEN] Funcionario encontrado: {funcionario['nombre_completo']}")

        # Obtener expedientes del funcionario
        cursor.execute("""
            SELECT
                e.nro_expediente,
                e.tipo_expediente,
                e.fecha_inicio,
                e.falta,
                e.estatus,
                e.decision
            FROM expedientes e
            WHERE e.funcionario_id = %s
            ORDER BY e.fecha_inicio ASC
        """, (funcionario['id'],))
        expedientes = cursor.fetchall()

        logger.info(f"[RECORD_GEN] Expedientes encontrados: {len(expedientes)}")

        funcionario_data = dict(funcionario)
        funcionario_data['expedientes_disciplinarios'] = [dict(exp) for exp in expedientes]

        return funcionario_data
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al obtener datos del funcionario {cedula}: {e}")
        raise HTTPException(status_code=500, detail=f"Error al obtener datos del funcionario: {e}")
    finally:
        if conn:
            conn.close()


def replace_static_placeholders(doc, data):
    """Reemplaza los marcadores de posición estáticos en el documento."""
    logger.info(f"[RECORD_GEN] Reemplazando marcadores estáticos")
    
    total_replacements = 0
    
    def replace_in_paragraphs(paragraphs):
        replacements = 0
        for paragraph in paragraphs:
            # Concatenar el texto de todos los runs para buscar el marcador completo
            full_text = "".join([run.text for run in paragraph.runs])
            
            # Logging especial para {NOMBRE}
            if "{NOMBRE}" in full_text:
                logger.info(f"[RECORD_GEN] ENCONTRADO {{NOMBRE}} en párrafo: '{full_text}'")
                logger.info(f"[RECORD_GEN] Runs del párrafo: {[run.text for run in paragraph.runs]}")
            
            # Verificar si hay algún marcador en este párrafo
            has_placeholder = False
            new_full_text = full_text
            
            for key, value in data.items():
                placeholder = f"{{{key}}}"
                if placeholder in new_full_text:
                    new_full_text = new_full_text.replace(placeholder, str(value))
                    logger.info(f"[RECORD_GEN] Reemplazado en párrafo: {placeholder} -> {value}")
                    has_placeholder = True
                    replacements += 1
            
            # Solo actualizar el párrafo si hubo reemplazos
            if has_placeholder:
                paragraph.clear()
                paragraph.add_run(new_full_text)
                
        return replacements

    # Reemplazar en párrafos normales
    paragraph_replacements = replace_in_paragraphs(doc.paragraphs)
    total_replacements += paragraph_replacements

    # Reemplazar en tablas
    table_replacements = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # También procesar párrafos dentro de las celdas
                cell_replacements = replace_in_paragraphs(cell.paragraphs)
                table_replacements += cell_replacements
                
                # Reemplazo directo en el texto de la celda como respaldo
                # Logging especial para {NOMBRE} en celdas
                if "{NOMBRE}" in cell.text:
                    logger.info(f"[RECORD_GEN] ENCONTRADO {{NOMBRE}} en celda: '{cell.text}'")
                
                for key, value in data.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        logger.info(f"[RECORD_GEN] Reemplazado en celda de tabla: {placeholder} -> {value}")
                        table_replacements += 1

    total_replacements += table_replacements
    
    logger.info(f"[RECORD_GEN] Total de reemplazos realizados: {total_replacements} (Párrafos: {paragraph_replacements}, Tablas: {table_replacements})")

def process_dynamic_table(doc, expedientes_data):
    """Procesa la tabla dinámica de expedientes de manera más eficiente y robusta."""
    if not expedientes_data:
        logger.info(f"[RECORD_GEN] No hay expedientes para procesar")
        return

    logger.info(f"[RECORD_GEN] Procesando tabla dinámica con {len(expedientes_data)} expedientes")

    # Marcadores que indican una fila de plantilla para expedientes
    dynamic_markers = ["{TIPO_EXP}", "{FECHA_EXP}", "{NRO_EXP}", "{FALTA}", "{STATUS}", "{DECISION}", "{1}"]
    
    for table_idx, table in enumerate(doc.tables):
        template_row_index = None
        template_row = None
        
        logger.debug(f"[RECORD_GEN] Analizando tabla {table_idx + 1}")
        
        # Buscar la fila que contiene los marcadores dinámicos
        for i, row in enumerate(table.rows):
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + " "
            
            # Verificar si esta fila contiene marcadores dinámicos
            markers_found = [marker for marker in dynamic_markers if marker in row_text]
            
            if markers_found:
                template_row_index = i
                template_row = row
                logger.info(f"[RECORD_GEN] Fila de plantilla encontrada en tabla {table_idx + 1}, fila {i}")
                logger.info(f"[RECORD_GEN] Marcadores encontrados: {markers_found}")
                break
        
        if template_row is None:
            logger.debug(f"[RECORD_GEN] No se encontraron marcadores dinámicos en tabla {table_idx + 1}")
            continue
        
        # Obtener la estructura de la fila plantilla y preservar el formato
        template_cells_data = []
        for cell in template_row.cells:
            cell_data = {
                'text': cell.text,
                'paragraphs': []
            }
            
            # Preservar información de párrafos y formato
            for paragraph in cell.paragraphs:
                para_data = {
                    'text': paragraph.text,
                    'runs': []
                }
                for run in paragraph.runs:
                    para_data['runs'].append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size
                    })
                cell_data['paragraphs'].append(para_data)
            
            template_cells_data.append(cell_data)
        
        logger.info(f"[RECORD_GEN] Estructura de plantilla capturada con {len(template_cells_data)} celdas")
        
        # Crear nuevas filas para cada expediente
        for idx, expediente in enumerate(expedientes_data):
            # Agregar nueva fila a la tabla
            new_row = table.add_row()
            
            # Mapear datos del expediente a los campos de la plantilla
            expediente_mapped = {
                "TIPO_EXP": expediente.get("tipo_expediente", ""),
                "FECHA_EXP": format_date(expediente.get("fecha_inicio")),
                "NRO_EXP": expediente.get("nro_expediente", ""),
                "FALTA": expediente.get("falta", ""),
                "STATUS": expediente.get("estatus", ""),
                "DECISION": expediente.get("decision", ""),
                "1": str(idx + 1)  # Enumerador
            }
            
            # Llenar cada celda de la nueva fila preservando el formato
            for j, cell in enumerate(new_row.cells):
                if j < len(template_cells_data):
                    cell_template = template_cells_data[j]
                    
                    # Limpiar la celda
                    cell.text = ""
                    
                    # Recrear párrafos con formato
                    for para_idx, para_template in enumerate(cell_template['paragraphs']):
                        if para_idx == 0:
                            # Usar el primer párrafo existente
                            paragraph = cell.paragraphs[0]
                        else:
                            # Crear nuevos párrafos
                            paragraph = cell.add_paragraph()
                        
                        # Limpiar párrafo
                        paragraph.clear()
                        
                        # Procesar runs con formato
                        for run_template in para_template['runs']:
                            run_text = run_template['text']
                            
                            # Reemplazar marcadores en el texto del run
                            for marker, value in expediente_mapped.items():
                                placeholder = f"{{{marker}}}"
                                if placeholder in run_text:
                                    run_text = run_text.replace(placeholder, str(value))
                            
                            # Crear nuevo run con formato
                            run = paragraph.add_run(run_text)
                            
                            # Aplicar formato
                            if run_template['bold']:
                                run.bold = True
                            if run_template['italic']:
                                run.italic = True
                            if run_template['underline']:
                                run.underline = True
                            if run_template['font_name']:
                                run.font.name = run_template['font_name']
                            if run_template['font_size']:
                                run.font.size = run_template['font_size']
            
            logger.debug(f"[RECORD_GEN] Expediente {idx + 1} agregado: {expediente_mapped}")
        
        # Limpiar la fila de plantilla original
        for cell in template_row.cells:
            cell.text = ""
        
        logger.info(f"[RECORD_GEN] Tabla dinámica procesada exitosamente")
        break  # Solo procesar la primera tabla que contenga marcadores

def format_date(date_value):
    """Formatea una fecha al formato DD/MM/YYYY."""
    if date_value is None:
        return ""
    
    if isinstance(date_value, str):
        try:
            # Intentar parsear diferentes formatos de fecha
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S"]:
                try:
                    parsed_date = datetime.strptime(date_value, fmt)
                    return parsed_date.strftime("%d/%m/%Y")
                except ValueError:
                    continue
            return date_value  # Si no se puede parsear, devolver como está
        except:
            return date_value
    
    if isinstance(date_value, (datetime, date)):
        return date_value.strftime("%d/%m/%Y")
    
    return str(date_value)

def validate_template_structure(doc):
    """Valida que la plantilla tenga la estructura esperada."""
    required_placeholders = [
        "{CEDULA}", "{NOMBRE}", "{JERARQUIA}", "{CANTIDAD_EXP}",
        "{HORA_ACTUAL}", "{FECHA_ACTUAL}", "{TIPO_SOLICITUD}",
        "{SUB_TOTAL_EXP}", "{TOTAL_EXP}"
    ]
    
    dynamic_placeholders = [
        "{TIPO_EXP}", "{FECHA_EXP}", "{NRO_EXP}", 
        "{FALTA}", "{STATUS}", "{DECISION}", "{1}"
    ]
    
    # Obtener texto de párrafos normales
    doc_text = ""
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + " "
    
    # Obtener texto de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text += cell.text + " "
    
    logger.debug(f"[RECORD_GEN] Texto total del documento: {doc_text[:500]}...")
    
    # Verificar marcadores faltantes
    missing_static = [p for p in required_placeholders if p not in doc_text]
    missing_dynamic = [p for p in dynamic_placeholders if p not in doc_text]
    
    if missing_static:
        logger.warning(f"[RECORD_GEN] Marcadores estáticos faltantes: {missing_static}")
    else:
        logger.info(f"[RECORD_GEN] Todos los marcadores estáticos encontrados")
    
    if missing_dynamic:
        logger.warning(f"[RECORD_GEN] Marcadores dinámicos faltantes: {missing_dynamic}")
    else:
        logger.info(f"[RECORD_GEN] Todos los marcadores dinámicos encontrados")
    
    # Mostrar marcadores encontrados
    found_static = [p for p in required_placeholders if p in doc_text]
    found_dynamic = [p for p in dynamic_placeholders if p in doc_text]
    
    if found_static:
        logger.info(f"[RECORD_GEN] Marcadores estáticos encontrados: {found_static}")
    if found_dynamic:
        logger.info(f"[RECORD_GEN] Marcadores dinámicos encontrados: {found_dynamic}")
    
    return len(missing_static) == 0 and len(missing_dynamic) == 0

@router.post("/documentos/generar-record-disciplinario")
async def generate_disciplinary_record(request: DocumentGenerationRequest):
    """
    Genera un record disciplinario en PDF a partir de una plantilla de Word
    y los datos de un funcionario y sus expedientes.
    """
    logger.info(f"[RECORD_GEN] Iniciando generación de record disciplinario")
    logger.info(f"[RECORD_GEN] Cédula: {request.cedula}, Tipo solicitud: {request.tipo_solicitud}")
    
    cedula = request.cedula
    tipo_solicitud = request.tipo_solicitud
    
    # Definir la ruta de la plantilla de Word
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'plantillas', 'Generar_records.docx')
    
    if not os.path.exists(template_path):
        logger.error(f"[RECORD_GEN] Plantilla no encontrada: {template_path}")
        raise HTTPException(status_code=404, detail=f"Plantilla de Word no encontrada en: {template_path}")

    # 1. Obtener datos del funcionario y sus expedientes
    funcionario_data = get_funcionario_data(cedula)
    if not funcionario_data:
        raise HTTPException(status_code=404, detail=f"Funcionario con cédula {cedula} no encontrado.")

    # 2. Preparar los datos para la plantilla
    current_time = datetime.now()
    expedientes = funcionario_data.get("expedientes_disciplinarios", [])
    
    # Obtener decisiones de expedientes para el campo DECISION
    decisiones = []
    for exp in expedientes:
        if exp.get("decision"):
            decisiones.append(exp.get("decision"))
    
    decision_text = "; ".join(decisiones) if decisiones else "SIN DECISIÓN REGISTRADA"
    
    # Datos estáticos para reemplazar (DECISION ahora es dinámico)
    static_data = {
        "CEDULA": funcionario_data.get("cedula", ""),
        "NOMBRE": funcionario_data.get("nombre_completo", ""),
        "JERARQUIA": funcionario_data.get("rango_actual", ""),
        "CANTIDAD_EXP": len(expedientes),
        "HORA_ACTUAL": current_time.strftime("%H:%M:%S"),
        "FECHA_ACTUAL": current_time.strftime("%d/%m/%Y"),
        "TIPO_SOLICITUD": tipo_solicitud,
        "SUB_TOTAL_EXP": len(expedientes),
        "TOTAL_EXP": len(expedientes)
    }
    
    logger.info(f"[RECORD_GEN] Datos estáticos preparados: {static_data}")

    # 3. Crear documento temporal modificado
    import tempfile
    import shutil
    
    try:
        # Crear directorio temporal
        temp_dir = tempfile.mkdtemp()
        temp_template_path = os.path.join(temp_dir, "modified_document.docx")
        
        # Copiar plantilla al directorio temporal
        shutil.copy2(template_path, temp_template_path)
        
        # Cargar la plantilla desde el archivo temporal
        doc = docx.Document(temp_template_path)
        logger.info(f"[RECORD_GEN] Documento modificado creado: {temp_template_path}")
        logger.info(f"[RECORD_GEN] Documento modificado cargado exitosamente")
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al crear documento temporal: {e}")
        raise HTTPException(status_code=500, detail=f"Error al crear documento temporal: {e}")

    # 4. Validar estructura de la plantilla
    template_valid = validate_template_structure(doc)
    if not template_valid:
        logger.warning(f"[RECORD_GEN] La plantilla tiene marcadores faltantes, pero continuando...")

    # 5. Reemplazar marcadores estáticos (incluyendo textboxes)
    replace_static_placeholders(doc, static_data)

    # 6. Procesar tabla dinámica de expedientes
    process_dynamic_table(doc, expedientes)

    # 7. Generar archivos de salida
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'temp_uploads')
    os.makedirs(output_dir, exist_ok=True)

    # Generar nombres únicos para los archivos
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = os.urandom(4).hex()
    base_filename = f"record_disciplinario_{cedula}_{timestamp}_{unique_id}"
    
    word_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"
    
    word_path = os.path.join(output_dir, word_filename)
    pdf_path = os.path.join(output_dir, pdf_filename)

    # 8. Guardar documento Word
    try:
        doc.save(word_path)
        logger.info(f"[RECORD_GEN] Documento Word guardado: {word_path}")
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al guardar documento Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error al guardar el documento de Word: {e}")

    # 9. Convertir a PDF
    try:
        convert(word_path, pdf_path)
        logger.info(f"[RECORD_GEN] Documento PDF generado: {pdf_path}")
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al convertir a PDF: {e}")
        logger.warning("[RECORD_GEN] Asegúrate de tener Microsoft Word o LibreOffice instalado")
        raise HTTPException(status_code=500, detail=f"Error al convertir a PDF: {e}")
    finally:
        # Limpiar archivo Word temporal
        if os.path.exists(word_path):
            try:
                os.remove(word_path)
                logger.info(f"[RECORD_GEN] Archivo Word temporal eliminado")
            except:
                logger.warning(f"[RECORD_GEN] No se pudo eliminar archivo Word temporal")

    # 10. Leer PDF y codificar en base64
    try:
        with open(pdf_path, "rb") as f:
            pdf_content = f.read()
        encoded_pdf = base64.b64encode(pdf_content).decode('utf-8')
        logger.info(f"[RECORD_GEN] PDF codificado en base64 exitosamente")
        
        return {
            "pdf_base64": encoded_pdf, 
            "message": "PDF de Record Disciplinario generado exitosamente.",
            "funcionario": {
                "cedula": funcionario_data.get("cedula", ""),
                "nombre": funcionario_data.get("nombre_completo", ""),
                "jerarquia": funcionario_data.get("rango_actual", ""),
                "total_expedientes": len(expedientes)
            }
        }
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al leer o codificar el PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error al procesar el PDF final: {e}")
    finally:
        # Limpiar archivo PDF temporal
        if os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
                logger.info(f"[RECORD_GEN] Archivo PDF temporal eliminado")
            except:
                logger.warning(f"[RECORD_GEN] No se pudo eliminar archivo PDF temporal")
        
        # Limpiar directorio temporal
        try:
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                logger.info(f"[RECORD_GEN] Directorio temporal eliminado")
        except:
            logger.warning(f"[RECORD_GEN] No se pudo eliminar directorio temporal")
