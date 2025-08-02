from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import docx
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx2pdf import convert
import os
import json
import base64
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime, date
import psycopg2
from psycopg2.extras import RealDictCursor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from lxml import etree

# Importar la configuración de la base de datos desde etl_app/config.py
try:
    from config import DATABASE_CONFIG as DB_CONFIG
except ImportError:
    # Fallback para desarrollo/testing si se ejecuta directamente el módulo
    DB_CONFIG = {
        "host": "localhost",
        "database": "jupe",
        "user": "postgres",
        "password": "12345678",
        "port": 5432
    }
    logging.warning("No se pudo importar DATABASE_CONFIG desde config.py. Usando configuración por defecto.")

router = APIRouter()
logger = logging.getLogger(__name__)

class DocumentGenerationRequest(BaseModel):
    cedula: str
    tipo_solicitud: str

def get_db_connection():
    """Crear conexión a la base de datos"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        logger.error(f"Error conectando a la base de datos: {e}")
        raise HTTPException(status_code=500, detail="Error de conexión a la base de datos")

def get_funcionario_data(cedula: str) -> Optional[Dict]:
    """Obtiene los datos del funcionario y sus expedientes de la base de datos."""
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)

        logger.info(f"[RECORD_GEN] Buscando funcionario con cédula: {cedula}")

        # Obtener datos del funcionario
        cursor.execute("""
            SELECT 
                f.id,
                f.cedula, 
                f.nombre_completo, 
                f.rango_actual
            FROM funcionarios f
            WHERE f.cedula = %s
        """, (cedula,))
        funcionario = cursor.fetchone()

        if not funcionario:
            logger.warning(f"[RECORD_GEN] Funcionario con cédula {cedula} no encontrado")
            return None

        logger.info(f"[RECORD_GEN] Funcionario encontrado: {funcionario['nombre_completo']}")

        # Obtener expedientes del funcionario
        cursor.execute("""
            SELECT
                e.nro_expediente,
                e.tipo_expediente,
                e.fecha_inicio,
                e.falta,
                e.estatus,
                e.decision
            FROM expedientes e
            WHERE e.funcionario_id = %s
            ORDER BY e.fecha_inicio ASC
        """, (funcionario['id'],))
        expedientes = cursor.fetchall()

        logger.info(f"[RECORD_GEN] Expedientes encontrados: {len(expedientes)}")

        funcionario_data = dict(funcionario)
        funcionario_data['expedientes_disciplinarios'] = [dict(exp) for exp in expedientes]

        return funcionario_data
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al obtener datos del funcionario {cedula}: {e}")
        raise HTTPException(status_code=500, detail=f"Error al obtener datos del funcionario: {e}")
    finally:
        if conn:
            conn.close()


def replace_static_placeholders(doc, data):
    """Reemplaza los marcadores de posición estáticos en el documento."""
    logger.info(f"[RECORD_GEN] Reemplazando marcadores estáticos")
    
    total_replacements = 0
    
    def replace_in_paragraphs(paragraphs):
        replacements = 0
        for paragraph in paragraphs:
            # Concatenar el texto de todos los runs para buscar el marcador completo
            full_text = "".join([run.text for run in paragraph.runs])
            
            # Logging especial para {NOMBRE}
            if "{NOMBRE}" in full_text:
                logger.info(f"[RECORD_GEN] ENCONTRADO {{NOMBRE}} en párrafo: '{full_text}'")
                logger.info(f"[RECORD_GEN] Runs del párrafo: {[run.text for run in paragraph.runs]}")
            
            # Verificar si hay algún marcador en este párrafo
            has_placeholder = False
            new_full_text = full_text
            
            for key, value in data.items():
                placeholder = f"{{{key}}}"
                if placeholder in new_full_text:
                    new_full_text = new_full_text.replace(placeholder, str(value))
                    logger.info(f"[RECORD_GEN] Reemplazado en párrafo: {placeholder} -> {value}")
                    has_placeholder = True
                    replacements += 1
            
            # Solo actualizar el párrafo si hubo reemplazos
            if has_placeholder:
                paragraph.clear()
                paragraph.add_run(new_full_text)
                
        return replacements

    # Reemplazar en párrafos normales
    paragraph_replacements = replace_in_paragraphs(doc.paragraphs)
    total_replacements += paragraph_replacements

    # Reemplazar en tablas
    table_replacements = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # También procesar párrafos dentro de las celdas
                cell_replacements = replace_in_paragraphs(cell.paragraphs)
                table_replacements += cell_replacements
                
                # Reemplazo directo en el texto de la celda como respaldo
                # Logging especial para {NOMBRE} en celdas
                if "{NOMBRE}" in cell.text:
                    logger.info(f"[RECORD_GEN] ENCONTRADO {{NOMBRE}} en celda: '{cell.text}'")
                
                for key, value in data.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        logger.info(f"[RECORD_GEN] Reemplazado en celda de tabla: {placeholder} -> {value}")
                        table_replacements += 1

    total_replacements += table_replacements
    
    logger.info(f"[RECORD_GEN] Total de reemplazos realizados: {total_replacements} (Párrafos: {paragraph_replacements}, Tablas: {table_replacements})")

def process_dynamic_table(doc, expedientes_data):
    """Procesa la tabla dinámica copiando la tabla completa para cada expediente."""
    if not expedientes_data:
        logger.info(f"[RECORD_GEN] No hay expedientes para procesar")
        return

    logger.info(f"[RECORD_GEN] Procesando tabla dinámica con {len(expedientes_data)} expedientes")

    # Marcadores que indican una tabla de plantilla para expedientes
    dynamic_markers = ["{TIPO_EXP}", "{FECHA_EXP}", "{NRO_EXP}", "{FALTA}", "{STATUS}", "{DECISION}", "{1}"]
    
    template_table = None
    template_table_element = None
    
    # Buscar la tabla que contiene los marcadores dinámicos
    for table_idx, table in enumerate(doc.tables):
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "
        
        # Verificar si esta tabla contiene marcadores dinámicos
        markers_found = [marker for marker in dynamic_markers if marker in table_text]
        
        if markers_found:
            template_table = table
            template_table_element = table._element
            logger.info(f"[RECORD_GEN] Tabla de plantilla encontrada: tabla {table_idx + 1}")
            logger.info(f"[RECORD_GEN] Marcadores encontrados: {markers_found}")
            break
    
    if template_table is None:
        logger.warning(f"[RECORD_GEN] No se encontró tabla con marcadores dinámicos")
        return
    
    # Crear contador manual para expedientes
    contador_expedientes = 0
    
    # Procesar cada expediente
    for idx, expediente in enumerate(expedientes_data):
        contador_expedientes += 1  # Incrementar contador
        logger.info(f"[RECORD_GEN] Procesando expediente {contador_expedientes}: {expediente.get('nro_expediente', 'N/A')}")
        
        # Mapear datos del expediente
        expediente_mapped = {
            "TIPO_EXP": expediente.get("tipo_expediente", ""),
            "FECHA_EXP": format_date(expediente.get("fecha_inicio")),
            "NRO_EXP": expediente.get("nro_expediente", ""),
            "FALTA": expediente.get("falta", ""),
            "STATUS": expediente.get("estatus", ""),
            "DECISION": expediente.get("decision", ""),
            "1": str(contador_expedientes)  # Enumerador manual (1, 2, 3, etc.)
        }
        
        logger.info(f"[RECORD_GEN] Datos del expediente {contador_expedientes}: {expediente_mapped}")
        logger.info(f"[RECORD_GEN] CONTADOR MANUAL: {{1}} = {str(contador_expedientes)}")
        
        if idx == 0:
            # Para el primer expediente, usar la tabla original
            current_table = template_table
        else:
            # Para expedientes adicionales, duplicar la tabla completa
            from copy import deepcopy
            
            # Crear una copia profunda del elemento de la tabla
            new_table_element = deepcopy(template_table_element)
            
            # Insertar la nueva tabla después de la tabla original con separación
            parent = template_table_element.getparent()
            table_index = list(parent).index(template_table_element)
            
            # Agregar párrafo de separación (0.5 cm)
            from docx.oxml.parser import parse_xml
            from docx.oxml.ns import nsdecls
            
            p_xml = f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="280"/></w:pPr></w:p>'  # 280 = aprox 0.5cm
            p_element = parse_xml(p_xml)
            
            # Insertar párrafo y tabla
            parent.insert(table_index + 1 + (idx - 1) * 2, p_element)
            parent.insert(table_index + 2 + (idx - 1) * 2, new_table_element)
            
            # Crear objeto Table de python-docx para la nueva tabla
            from docx.table import Table
            current_table = Table(new_table_element, doc)
        
        # Reemplazar marcadores en la tabla actual
        for row in current_table.rows:
            for cell in row.cells:
                # Primero, reemplazo directo en el texto de la celda para {1}
                if "{1}" in cell.text:
                    cell.text = cell.text.replace("{1}", str(contador_expedientes))
                    logger.info(f"[RECORD_GEN] ENUMERACIÓN DIRECTA: Expediente {contador_expedientes} - Reemplazado {{1}} con {str(contador_expedientes)}")
                
                # Procesar párrafos dentro de la celda
                for paragraph in cell.paragraphs:
                    full_text = "".join([run.text for run in paragraph.runs])
                    
                    # Verificar si hay marcadores en este párrafo
                    has_placeholder = False
                    new_full_text = full_text
                    
                    for marker, value in expediente_mapped.items():
                        placeholder = f"{{{marker}}}"
                        if placeholder in new_full_text:
                            new_full_text = new_full_text.replace(placeholder, str(value))
                            logger.info(f"[RECORD_GEN] Reemplazado en tabla expediente {contador_expedientes}: {placeholder} -> {value}")
                            has_placeholder = True
                            
                            # Logging especial para {1}
                            if marker == "1":
                                logger.info(f"[RECORD_GEN] ENUMERACIÓN: Expediente {contador_expedientes} - Reemplazando {{1}} con {value}")
                    
                    # Solo actualizar el párrafo si hubo reemplazos
                    if has_placeholder:
                        paragraph.clear()
                        run = paragraph.add_run(new_full_text)
                        # Aplicar formato: Calibri, negrita, tamaño 6
                        run.font.name = 'Calibri'
                        run.font.bold = True
                        run.font.size = Pt(6)
        
        logger.info(f"[RECORD_GEN] Expediente {contador_expedientes} procesado exitosamente")
    
    logger.info(f"[RECORD_GEN] Tabla dinámica procesada exitosamente para {len(expedientes_data)} expedientes")

def format_date(date_value):
    """Formatea una fecha al formato DD/MM/YYYY."""
    if date_value is None:
        return ""
    
    if isinstance(date_value, str):
        try:
            # Intentar parsear diferentes formatos de fecha
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S"]:
                try:
                    parsed_date = datetime.strptime(date_value, fmt)
                    return parsed_date.strftime("%d/%m/%Y")
                except ValueError:
                    continue
            return date_value  # Si no se puede parsear, devolver como está
        except:
            return date_value
    
    if isinstance(date_value, (datetime, date)):
        return date_value.strftime("%d/%m/%Y")
    
    return str(date_value)

def validate_template_structure(doc):
    """Valida que la plantilla tenga la estructura esperada."""
    required_placeholders = [
        "{CEDULA}", "{NOMBRE}", "{JERARQUIA}", "{CANTIDAD_EXP}",
        "{HORA_ACTUAL}", "{FECHA_ACTUAL}", "{TIPO_SOLICITUD}",
        "{SUB_TOTAL_EXP}", "{TOTAL_EXP}"
    ]
    
    dynamic_placeholders = [
        "{TIPO_EXP}", "{FECHA_EXP}", "{NRO_EXP}", 
        "{FALTA}", "{STATUS}", "{DECISION}", "{1}"
    ]
    
    # Obtener texto de párrafos normales
    doc_text = ""
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + " "
    
    # Obtener texto de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text += cell.text + " "
    
    logger.debug(f"[RECORD_GEN] Texto total del documento: {doc_text[:500]}...")
    
    # Verificar marcadores faltantes
    missing_static = [p for p in required_placeholders if p not in doc_text]
    missing_dynamic = [p for p in dynamic_placeholders if p not in doc_text]
    
    if missing_static:
        logger.warning(f"[RECORD_GEN] Marcadores estáticos faltantes: {missing_static}")
    else:
        logger.info(f"[RECORD_GEN] Todos los marcadores estáticos encontrados")
    
    if missing_dynamic:
        logger.warning(f"[RECORD_GEN] Marcadores dinámicos faltantes: {missing_dynamic}")
    else:
        logger.info(f"[RECORD_GEN] Todos los marcadores dinámicos encontrados")
    
    # Mostrar marcadores encontrados
    found_static = [p for p in required_placeholders if p in doc_text]
    found_dynamic = [p for p in dynamic_placeholders if p in doc_text]
    
    if found_static:
        logger.info(f"[RECORD_GEN] Marcadores estáticos encontrados: {found_static}")
    if found_dynamic:
        logger.info(f"[RECORD_GEN] Marcadores dinámicos encontrados: {found_dynamic}")
    
    return len(missing_static) == 0 and len(missing_dynamic) == 0

@router.post("/documentos/generar-record-disciplinario")
async def generate_disciplinary_record(request: DocumentGenerationRequest):
    """
    Genera un record disciplinario en PDF a partir de una plantilla de Word
    y los datos de un funcionario y sus expedientes.
    """
    logger.info(f"[RECORD_GEN] Iniciando generación de record disciplinario")
    logger.info(f"[RECORD_GEN] Cédula: {request.cedula}, Tipo solicitud: {request.tipo_solicitud}")
    
    cedula = request.cedula
    tipo_solicitud = request.tipo_solicitud
    
    # Definir la ruta de la plantilla de Word
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'plantillas', 'Generar_records.docx')
    
    if not os.path.exists(template_path):
        logger.error(f"[RECORD_GEN] Plantilla no encontrada: {template_path}")
        raise HTTPException(status_code=404, detail=f"Plantilla de Word no encontrada en: {template_path}")

    # 1. Obtener datos del funcionario y sus expedientes
    funcionario_data = get_funcionario_data(cedula)
    if not funcionario_data:
        raise HTTPException(status_code=404, detail=f"Funcionario con cédula {cedula} no encontrado.")

    # 2. Preparar los datos para la plantilla
    current_time = datetime.now()
    expedientes = funcionario_data.get("expedientes_disciplinarios", [])
    
    # Obtener decisiones de expedientes para el campo DECISION
    decisiones = []
    for exp in expedientes:
        if exp.get("decision"):
            decisiones.append(exp.get("decision"))
    
    decision_text = "; ".join(decisiones) if decisiones else "SIN DECISIÓN REGISTRADA"
    
    # Datos estáticos para reemplazar (DECISION ahora es dinámico)
    static_data = {
        "CEDULA": funcionario_data.get("cedula", ""),
        "NOMBRE": funcionario_data.get("nombre_completo", ""),
        "JERARQUIA": funcionario_data.get("rango_actual", ""),
        "CANTIDAD_EXP": len(expedientes),
        "HORA_ACTUAL": current_time.strftime("%H:%M:%S"),
        "FECHA_ACTUAL": current_time.strftime("%d/%m/%Y"),
        "TIPO_SOLICITUD": tipo_solicitud,
        "SUB_TOTAL_EXP": len(expedientes),
        "TOTAL_EXP": len(expedientes)
    }
    
    logger.info(f"[RECORD_GEN] Datos estáticos preparados: {static_data}")

    # 3. Crear documento temporal modificado
    import tempfile
    import shutil
    
    try:
        # Crear directorio temporal
        temp_dir = tempfile.mkdtemp()
        temp_template_path = os.path.join(temp_dir, "modified_document.docx")
        
        # Copiar plantilla al directorio temporal
        shutil.copy2(template_path, temp_template_path)
        
        # Cargar la plantilla desde el archivo temporal
        doc = docx.Document(temp_template_path)
        logger.info(f"[RECORD_GEN] Documento modificado creado: {temp_template_path}")
        logger.info(f"[RECORD_GEN] Documento modificado cargado exitosamente")
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al crear documento temporal: {e}")
        raise HTTPException(status_code=500, detail=f"Error al crear documento temporal: {e}")

    # 4. Validar estructura de la plantilla
    template_valid = validate_template_structure(doc)
    if not template_valid:
        logger.warning(f"[RECORD_GEN] La plantilla tiene marcadores faltantes, pero continuando...")

    # 5. Reemplazar marcadores estáticos (incluyendo textboxes)
    replace_static_placeholders(doc, static_data)

    # 6. Procesar tabla dinámica de expedientes
    process_dynamic_table(doc, expedientes)

    # 7. Generar archivos de salida
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'temp_uploads')
    os.makedirs(output_dir, exist_ok=True)

    # Generar nombres únicos para los archivos
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = os.urandom(4).hex()
    base_filename = f"record_disciplinario_{cedula}_{timestamp}_{unique_id}"
    
    word_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"
    
    word_path = os.path.join(output_dir, word_filename)
    pdf_path = os.path.join(output_dir, pdf_filename)

    # 8. Guardar documento Word
    try:
        doc.save(word_path)
        logger.info(f"[RECORD_GEN] Documento Word guardado: {word_path}")
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al guardar documento Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error al guardar el documento de Word: {e}")

    # 9. Convertir a PDF con inicialización COM y fallback
    pdf_success = False
    try:
        logger.info(f"[RECORD_GEN] Iniciando conversión a PDF...")
        
        # Inicializar COM para Windows
        try:
            import pythoncom
            pythoncom.CoInitialize()
            logger.info(f"[RECORD_GEN] COM inicializado para Windows")
        except ImportError:
            logger.info(f"[RECORD_GEN] pythoncom no disponible, continuando sin inicialización COM")
        except Exception as com_error:
            logger.warning(f"[RECORD_GEN] Error al inicializar COM: {com_error}")
        
        try:
            import threading
            import time
            import subprocess
            
            logger.info(f"[RECORD_GEN] Verificando archivos antes de conversión...")
            logger.info(f"[RECORD_GEN] Archivo Word existe: {os.path.exists(word_path)}")
            logger.info(f"[RECORD_GEN] Ruta Word: {word_path}")
            logger.info(f"[RECORD_GEN] Ruta PDF destino: {pdf_path}")
            
            # Verificar tamaño del archivo Word
            try:
                word_size = os.path.getsize(word_path)
                logger.info(f"[RECORD_GEN] Tamaño del archivo Word: {word_size} bytes")
                if word_size == 0:
                    logger.error(f"[RECORD_GEN] PROBLEMA: El archivo Word está vacío!")
                elif word_size < 1000:
                    logger.warning(f"[RECORD_GEN] ADVERTENCIA: El archivo Word es muy pequeño ({word_size} bytes)")
            except Exception as e:
                logger.error(f"[RECORD_GEN] Error al verificar tamaño del archivo Word: {e}")
            
            # Variable para controlar el resultado de la conversión
            conversion_result = {"success": False, "error": None, "method": None}
            
            def convert_with_docx2pdf():
                try:
                    logger.info(f"[RECORD_GEN] Intentando conversión con docx2pdf...")
                    
                    # Reinicializar COM específicamente para este hilo
                    try:
                        import pythoncom
                        pythoncom.CoInitialize()
                        logger.info(f"[RECORD_GEN] COM reinicializado en hilo de conversión")
                    except Exception as com_error:
                        logger.warning(f"[RECORD_GEN] Error al reinicializar COM en hilo: {com_error}")
                    
                    convert(word_path, pdf_path)
                    conversion_result["success"] = True
                    conversion_result["method"] = "docx2pdf"
                    logger.info(f"[RECORD_GEN] Documento PDF generado con docx2pdf: {pdf_path}")
                    
                    # Limpiar COM del hilo
                    try:
                        pythoncom.CoUninitialize()
                    except:
                        pass
                        
                except Exception as e:
                    conversion_result["error"] = e
                    logger.error(f"[RECORD_GEN] Error en docx2pdf: {e}")
                    
                    # Limpiar COM en caso de error
                    try:
                        pythoncom.CoUninitialize()
                    except:
                        pass
            
            def convert_with_libreoffice():
                try:
                    logger.info(f"[RECORD_GEN] Intentando conversión con LibreOffice...")
                    
                    # Buscar LibreOffice en diferentes ubicaciones comunes
                    libreoffice_paths = [
                        "libreoffice",  # En PATH
                        "soffice",      # Comando alternativo
                        r"C:\Program Files\LibreOffice\program\soffice.exe",
                        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                        r"C:\Users\{}\AppData\Local\Programs\LibreOffice\program\soffice.exe".format(os.getenv('USERNAME', '')),
                    ]
                    
                    libreoffice_cmd = None
                    for path in libreoffice_paths:
                        try:
                            logger.info(f"[RECORD_GEN] Probando LibreOffice en: {path}")
                            
                            # Para comandos en PATH, probar directamente
                            if path in ["libreoffice", "soffice"]:
                                test_result = subprocess.run([path, "--version"], 
                                                           capture_output=True, timeout=5, text=True)
                                logger.info(f"[RECORD_GEN] Resultado de {path} --version: returncode={test_result.returncode}")
                                logger.info(f"[RECORD_GEN] stdout: {test_result.stdout[:100]}...")
                                logger.info(f"[RECORD_GEN] stderr: {test_result.stderr[:100]}...")
                                
                                if test_result.returncode == 0:
                                    libreoffice_cmd = path
                                    logger.info(f"[RECORD_GEN] LibreOffice encontrado en: {path}")
                                    break
                            
                            # Para rutas de archivos, verificar que existe primero
                            elif os.path.exists(path):
                                test_result = subprocess.run([path, "--version"], 
                                                           capture_output=True, timeout=5, text=True)
                                logger.info(f"[RECORD_GEN] Resultado de {path} --version: returncode={test_result.returncode}")
                                
                                if test_result.returncode == 0:
                                    libreoffice_cmd = path
                                    logger.info(f"[RECORD_GEN] LibreOffice encontrado en: {path}")
                                    break
                            else:
                                logger.debug(f"[RECORD_GEN] Archivo no existe: {path}")
                                
                        except (subprocess.TimeoutExpired, FileNotFoundError, OSError) as e:
                            logger.info(f"[RECORD_GEN] Error probando {path}: {e}")
                            continue
                    
                    if not libreoffice_cmd:
                        logger.warning(f"[RECORD_GEN] LibreOffice no detectado con --version, intentando conversión directa...")
                        # Intentar directamente con soffice ya que sabemos que funciona
                        libreoffice_cmd = "soffice"
                    
                    # Ejecutar conversión con PATH del sistema
                    logger.info(f"[RECORD_GEN] Ejecutando conversión con: {libreoffice_cmd}")
                    
                    # Usar shell=True para acceder al PATH del sistema
                    cmd_str = f'"{libreoffice_cmd}" --headless --convert-to pdf --outdir "{os.path.dirname(pdf_path)}" "{word_path}"'
                    logger.info(f"[RECORD_GEN] Comando completo: {cmd_str}")
                    
                    result = subprocess.run(cmd_str, shell=True, timeout=30, capture_output=True, text=True)
                    
                    logger.info(f"[RECORD_GEN] LibreOffice returncode: {result.returncode}")
                    logger.info(f"[RECORD_GEN] LibreOffice stdout: {result.stdout}")
                    logger.info(f"[RECORD_GEN] LibreOffice stderr: {result.stderr}")
                    
                    # Verificar si el comando fue exitoso
                    if result.returncode == 0:
                        # LibreOffice puede generar el PDF con nombre basado en el archivo Word
                        word_basename = os.path.splitext(os.path.basename(word_path))[0]
                        libreoffice_pdf_path = os.path.join(os.path.dirname(pdf_path), f"{word_basename}.pdf")
                        
                        logger.info(f"[RECORD_GEN] Verificando PDF en ruta esperada: {pdf_path}")
                        logger.info(f"[RECORD_GEN] Verificando PDF en ruta LibreOffice: {libreoffice_pdf_path}")
                        
                        if os.path.exists(pdf_path):
                            conversion_result["success"] = True
                            conversion_result["method"] = "libreoffice"
                            logger.info(f"[RECORD_GEN] Documento PDF generado con LibreOffice: {pdf_path}")
                        elif os.path.exists(libreoffice_pdf_path):
                            # Mover el archivo a la ubicación esperada
                            import shutil
                            shutil.move(libreoffice_pdf_path, pdf_path)
                            conversion_result["success"] = True
                            conversion_result["method"] = "libreoffice"
                            logger.info(f"[RECORD_GEN] Documento PDF generado con LibreOffice y movido: {pdf_path}")
                        else:
                            logger.error(f"[RECORD_GEN] LibreOffice no generó el archivo PDF en ninguna ubicación")
                            # Listar archivos en el directorio para debug
                            try:
                                files_in_dir = os.listdir(os.path.dirname(pdf_path))
                                logger.info(f"[RECORD_GEN] Archivos en directorio: {files_in_dir}")
                            except:
                                pass
                    else:
                        logger.error(f"[RECORD_GEN] LibreOffice falló con código de retorno: {result.returncode}")
                        
                except subprocess.CalledProcessError as e:
                    logger.error(f"[RECORD_GEN] Error en LibreOffice (CalledProcessError): {e}")
                    logger.error(f"[RECORD_GEN] LibreOffice stdout: {e.stdout}")
                    logger.error(f"[RECORD_GEN] LibreOffice stderr: {e.stderr}")
                except subprocess.TimeoutExpired as e:
                    logger.error(f"[RECORD_GEN] Timeout en LibreOffice: {e}")
                except FileNotFoundError as e:
                    logger.error(f"[RECORD_GEN] LibreOffice no encontrado: {e}")
                except Exception as e:
                    logger.error(f"[RECORD_GEN] Error general en LibreOffice: {e}")
            
            # Intentar primero con docx2pdf
            logger.info(f"[RECORD_GEN] Iniciando conversión con docx2pdf (timeout 30s)...")
            conversion_thread = threading.Thread(target=convert_with_docx2pdf)
            conversion_thread.daemon = True
            conversion_thread.start()
            conversion_thread.join(timeout=30)
            
            if conversion_thread.is_alive():
                logger.error(f"[RECORD_GEN] Timeout en docx2pdf (30 segundos)")
                logger.info(f"[RECORD_GEN] Intentando con LibreOffice como alternativa...")
                
                # Intentar con LibreOffice
                convert_with_libreoffice()
                
            elif conversion_result["success"]:
                pdf_success = True
                logger.info(f"[RECORD_GEN] Conversión exitosa con {conversion_result['method']}")
            else:
                logger.error(f"[RECORD_GEN] Error en docx2pdf: {conversion_result.get('error', 'Error desconocido')}")
                logger.info(f"[RECORD_GEN] Intentando con LibreOffice como alternativa...")
                
                # Intentar con LibreOffice
                convert_with_libreoffice()
                
                if conversion_result["success"]:
                    pdf_success = True
                    logger.info(f"[RECORD_GEN] Conversión exitosa con {conversion_result['method']}")
                else:
                    logger.error(f"[RECORD_GEN] Falló tanto docx2pdf como LibreOffice")
                    logger.info(f"[RECORD_GEN] Continuando con documento Word como alternativa")
            
            # Verificar una vez más el estado final
            logger.info(f"[RECORD_GEN] Estado final de conversión: pdf_success={pdf_success}, conversion_result={conversion_result}")
                
        except Exception as e:
            logger.error(f"[RECORD_GEN] Error general en conversión a PDF: {e}")
            logger.info(f"[RECORD_GEN] Continuando con documento Word como alternativa")
        
        # Limpiar COM
        try:
            pythoncom.CoUninitialize()
        except:
            pass
            
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error general en conversión a PDF: {e}")
    finally:
        # NO eliminar archivo Word temporal aquí, lo necesitamos para el fallback
        pass

    # 10. Procesar resultado final (PDF o Word como fallback)
    try:
        logger.info(f"[RECORD_GEN] Procesando resultado final...")
        logger.info(f"[RECORD_GEN] pdf_success = {pdf_success}")
        logger.info(f"[RECORD_GEN] PDF existe en {pdf_path}: {os.path.exists(pdf_path)}")
        
        if pdf_success and os.path.exists(pdf_path):
            # Si el PDF se generó exitosamente
            pdf_size = os.path.getsize(pdf_path)
            logger.info(f"[RECORD_GEN] Tamaño del PDF: {pdf_size} bytes")
            
            with open(pdf_path, "rb") as f:
                pdf_content = f.read()
            encoded_pdf = base64.b64encode(pdf_content).decode('utf-8')
            logger.info(f"[RECORD_GEN] PDF codificado en base64 exitosamente")
            
            return {
                "pdf_base64": encoded_pdf,
                "word_base64": None,
                "message": "PDF de Record Disciplinario generado exitosamente.",
                "funcionario": {
                    "cedula": funcionario_data.get("cedula", ""),
                    "nombre": funcionario_data.get("nombre_completo", ""),
                    "jerarquia": funcionario_data.get("rango_actual", ""),
                    "total_expedientes": len(expedientes)
                }
            }
        else:
            # Si falló la conversión a PDF, devolver documento Word
            logger.info(f"[RECORD_GEN] Devolviendo documento Word como alternativa")
            logger.info(f"[RECORD_GEN] Razón: pdf_success={pdf_success}, archivo_existe={os.path.exists(pdf_path)}")
            
            with open(word_path, "rb") as f:
                word_content = f.read()
            encoded_word = base64.b64encode(word_content).decode('utf-8')
            
            return {
                "pdf_base64": None,
                "word_base64": encoded_word,
                "message": "Documento Word generado exitosamente. No se pudo convertir a PDF.",
                "funcionario": {
                    "cedula": funcionario_data.get("cedula", ""),
                    "nombre": funcionario_data.get("nombre_completo", ""),
                    "jerarquia": funcionario_data.get("rango_actual", ""),
                    "total_expedientes": len(expedientes)
                }
            }
            
    except Exception as e:
        logger.error(f"[RECORD_GEN] Error al procesar documento final: {e}")
        raise HTTPException(status_code=500, detail=f"Error al procesar documento final: {e}")
    finally:
        # Limpiar archivos temporales
        if os.path.exists(word_path):
            try:
                os.remove(word_path)
                logger.info(f"[RECORD_GEN] Archivo Word temporal eliminado")
            except:
                logger.warning(f"[RECORD_GEN] No se pudo eliminar archivo Word temporal")
        
        if os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
                logger.info(f"[RECORD_GEN] Archivo PDF temporal eliminado")
            except:
                logger.warning(f"[RECORD_GEN] No se pudo eliminar archivo PDF temporal")
        
        # Limpiar directorio temporal
        try:
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                logger.info(f"[RECORD_GEN] Directorio temporal eliminado")
        except:
            logger.warning(f"[RECORD_GEN] No se pudo eliminar directorio temporal")
